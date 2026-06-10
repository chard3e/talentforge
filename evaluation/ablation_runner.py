"""
TalentForge Ablation Study Runner

Desteklenen modeller:
  HuggingFace router : Qwen/Qwen2.5-7B-Instruct, Qwen/Qwen2.5-72B-Instruct
                       meta-llama/Llama-3.1-8B-Instruct
                       mistralai/Mixtral-8x7B-Instruct-v0.1
  OpenAI             : gpt-4o-mini, gpt-4o
  Google AI Studio   : gemini-1.5-flash, gemini-1.5-pro
  Groq               : llama-3.1-8b-instant, mixtral-8x7b-32768

Gerekli .env değişkenleri:
  HF_TOKEN, OPENAI_API_KEY, GOOGLE_API_KEY, GROQ_API_KEY

Kullanım:
  uv run python evaluation/ablation_runner.py \\
    --cv_dir data/cvs --n_cvs 20 \\
    --configs SYS-B \\
    --models gpt-4o-mini gemini-1.5-flash Qwen/Qwen2.5-7B-Instruct
"""

import argparse
import json
import logging
import os
import re
import sys
import time
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from dotenv import load_dotenv
load_dotenv()

sys.path.insert(0, str(Path(__file__).parent.parent))

from evaluation.triple_evaluator import make_triples

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")

LLM_MAX_TOKENS = int(os.getenv("EVAL_LLM_MAX_TOKENS", "8192"))

# ── Model → endpoint yönlendirme ─────────────────────────────────────

def _get_client(model: str):
    """Model adına göre doğru OpenAI-uyumlu client döndürür."""
    from openai import OpenAI

    if model.startswith("gpt"):
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY bulunamadı")
        return OpenAI(api_key=api_key), model

    if model.startswith("gemini"):
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key:
            raise ValueError("GOOGLE_API_KEY bulunamadı")
        return OpenAI(
            base_url="https://generativelanguage.googleapis.com/v1beta/openai/",
            api_key=api_key,
        ), model

    if model.startswith("llama") or model.startswith("mixtral"):
        # Kısa isimler → Groq
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            raise ValueError("GROQ_API_KEY bulunamadı")
        return OpenAI(
            base_url="https://api.groq.com/openai/v1",
            api_key=api_key,
        ), model

    # HuggingFace router (Qwen, meta-llama, mistralai/Mixtral vb.)
    api_key = os.getenv("HF_TOKEN")
    if not api_key:
        raise ValueError("HF_TOKEN bulunamadı")
    return OpenAI(
        base_url="https://router.huggingface.co/v1",
        api_key=api_key,
    ), model


# ── LLM çağrısı ───────────────────────────────────────────────────────

def call_llm(system_prompt: str, user_prompt: str, model: str) -> Optional[Dict]:
    try:
        client, model_id = _get_client(model)
    except ValueError as e:
        logger.warning(f"  API key hatası: {e}")
        return None

    try:
        response = client.chat.completions.create(
            model=model_id,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0,
            max_tokens=LLM_MAX_TOKENS,
        )
        raw_text = response.choices[0].message.content or ""
        result = _parse_json_response(raw_text)
        if result is None:
            logger.warning(f"    ⚠️ JSON parse başarısız. İlk 200: {raw_text[:200]!r}")
        return result
    except Exception as e:
        logger.warning(f"  LLM hatası ({model}): {e}")
        return None


def _parse_json_response(text: str) -> Optional[Dict]:
    text = text.strip()
    fenced = re.search(r"```(?:json)?\s*(.*?)\s*```", text, re.DOTALL | re.IGNORECASE)
    candidates = []
    if fenced:
        candidates.append(fenced.group(1).strip())
    candidates.append(re.sub(r"```(?:json)?", "", text, flags=re.IGNORECASE).strip().strip("`").strip())

    decoder = json.JSONDecoder()
    for candidate in candidates:
        try:
            return json.loads(candidate)
        except Exception:
            pass
        brace_index = candidate.find("{")
        if brace_index >= 0:
            try:
                parsed, _ = decoder.raw_decode(candidate[brace_index:])
                if isinstance(parsed, dict):
                    return parsed
            except Exception:
                pass

    # Last resort: trim trailing prose after the final closing brace.
    for candidate in candidates:
        start = candidate.find("{")
        end = candidate.rfind("}")
        if start >= 0 and end > start:
            try:
                return json.loads(candidate[start:end + 1])
            except Exception:
                pass
    return None


# ── Prompt varyantları ────────────────────────────────────────────────

def get_prompt_config(config_name: str):
    sys.path.insert(0, str(Path(__file__).parent))
    if config_name == "BL-1":
        from prompts_bl1 import SYSTEM_PROMPT_BL1, build_user_prompt_bl1
        return SYSTEM_PROMPT_BL1, build_user_prompt_bl1
    elif config_name == "BL-2":
        from prompts_bl2 import SYSTEM_PROMPT_BL2, build_user_prompt_bl2
        return SYSTEM_PROMPT_BL2, build_user_prompt_bl2
    elif config_name == "SYS-A":
        from prompts_sysA import SYSTEM_PROMPT_SYSA, build_user_prompt_sysA
        return SYSTEM_PROMPT_SYSA, build_user_prompt_sysA
    elif config_name in ("SYS-B", "SYS-C"):
        from app.extraction.prompts import SYSTEM_PROMPT, build_user_prompt
        return SYSTEM_PROMPT, build_user_prompt
    raise ValueError(f"Bilinmeyen konfigürasyon: {config_name}")


# ── Metrik fonksiyonları ──────────────────────────────────────────────

def normalize(s: str) -> str:
    if s is None:
        return ""
    return re.sub(r"\s+", " ", str(s).lower().strip())


def _as_list(value) -> List:
    if value is None:
        return []
    return value if isinstance(value, list) else [value]


def set_f1(pred_items: List, gold_items: List) -> float:
    p = {normalize(x) for x in pred_items if normalize(x)}
    g = {normalize(x) for x in gold_items if normalize(x)}
    if not p and not g:
        return 1.0
    if not p or not g:
        return 0.0
    tp = len(p & g)
    prec = tp / len(p)
    rec = tp / len(g)
    return 0.0 if prec + rec == 0 else 2 * prec * rec / (prec + rec)


def compute_ner_f1(pred: Dict, gold: Dict) -> Dict[str, float]:
    pred_skills = [s.get("name", s) if isinstance(s, dict) else s for s in _as_list(pred.get("skills"))]
    gold_skills = [s.get("name", s) if isinstance(s, dict) else s for s in _as_list(gold.get("skills"))]
    skill_f1 = set_f1(pred_skills, gold_skills)

    pred_companies = [e.get("company_name", "") for e in _as_list(pred.get("experiences")) if isinstance(e, dict)]
    gold_companies = [e.get("company_name", "") for e in _as_list(gold.get("experiences")) if isinstance(e, dict)]
    company_f1 = set_f1(pred_companies, gold_companies)

    pred_edu = [e.get("institution", "") for e in _as_list(pred.get("educations")) if isinstance(e, dict)]
    gold_edu = [e.get("institution", "") for e in _as_list(gold.get("educations")) if isinstance(e, dict)]
    edu_f1 = set_f1(pred_edu, gold_edu)

    return {
        "skill_f1":     round(skill_f1, 3),
        "company_f1":   round(company_f1, 3),
        "education_f1": round(edu_f1, 3),
        "overall_f1":   round((skill_f1 + company_f1 + edu_f1) / 3, 3),
    }


def compute_re_f1(pred: Dict, gold: Dict) -> float:
    gold_exps = _as_list(gold.get("experiences"))
    if not gold_exps:
        return 1.0

    gold_by_company: Dict[str, List[str]] = {}
    for exp in gold_exps:
        key = normalize(exp.get("company_name", ""))
        if key:
            gold_by_company[key] = exp.get("skills_used", [])

    if not gold_by_company:
        return 0.0

    f1_scores = []
    for gold_company, gold_skills in gold_by_company.items():
        best_f1 = 0.0
        for pred_exp in _as_list(pred.get("experiences")):
            if not isinstance(pred_exp, dict):
                continue
            pred_company = normalize(pred_exp.get("company_name", ""))
            if (pred_company == gold_company
                    or gold_company in pred_company
                    or pred_company in gold_company):
                f1 = set_f1(_as_list(pred_exp.get("skills_used")), _as_list(gold_skills))
                best_f1 = max(best_f1, f1)
        f1_scores.append(best_f1)

    return round(sum(f1_scores) / len(f1_scores), 3) if f1_scores else 0.0


def compute_kg_triple_metrics(pred: Dict, gold: Dict) -> Dict[str, float]:
    pred_triples = make_triples(pred)
    gold_triples = make_triples(gold)
    if not pred_triples and not gold_triples:
        return {
            "precision": 1.0,
            "recall": 1.0,
            "f1": 1.0,
            "tp": 0,
            "pred_count": 0,
            "gold_count": 0,
        }
    if not pred_triples or not gold_triples:
        return {
            "precision": 0.0,
            "recall": 0.0,
            "f1": 0.0,
            "tp": 0,
            "pred_count": len(pred_triples),
            "gold_count": len(gold_triples),
        }

    tp = len(pred_triples & gold_triples)
    precision = tp / len(pred_triples)
    recall = tp / len(gold_triples)
    f1 = 0.0 if precision + recall == 0 else 2 * precision * recall / (precision + recall)
    return {
        "precision": round(precision, 3),
        "recall": round(recall, 3),
        "f1": round(f1, 3),
        "tp": tp,
        "pred_count": len(pred_triples),
        "gold_count": len(gold_triples),
    }


EVIDENCE_WORD_THRESHOLD = 0.60
FUZZY_MATCH_THRESHOLD = 0.50
CONFIDENCE_MIN_THRESHOLD = 0.60


def _words(text: str) -> set:
    return set(re.findall(r"\b[a-z0-9+#./-]{3,}\b", normalize(text)))


def _word_overlap_ratio(short_text: str, long_text: str) -> float:
    short_words = _words(short_text)
    if not short_words:
        return 0.0
    long_words = _words(long_text)
    return len(short_words & long_words) / len(short_words)


def _phrase_supported(phrase: str, cv_text: str, threshold: float = FUZZY_MATCH_THRESHOLD) -> bool:
    phrase_norm = normalize(phrase)
    cv_norm = normalize(cv_text)
    if not phrase_norm:
        return True
    if phrase_norm in cv_norm:
        return True
    return _word_overlap_ratio(phrase_norm, cv_norm) >= threshold


def _get_confidence(item: Dict) -> Optional[float]:
    value = item.get("confidence") if isinstance(item, dict) else None
    try:
        return float(value) if value is not None else None
    except (TypeError, ValueError):
        return None


def _evidence_status(value: str, evidence: str, cv_text: str, confidence: Optional[float] = None) -> str:
    if confidence is not None and confidence < CONFIDENCE_MIN_THRESHOLD:
        return "quarantine"

    evidence = normalize(evidence)
    if evidence:
        overlap = _word_overlap_ratio(evidence, cv_text)
        if evidence in normalize(cv_text) or overlap >= EVIDENCE_WORD_THRESHOLD:
            return "supported"
        if overlap >= FUZZY_MATCH_THRESHOLD:
            return "warning"
        return "quarantine"

    if _phrase_supported(value, cv_text):
        return "supported"
    return "quarantine"


def _count_status(metrics: Dict, status: str, bucket: str):
    metrics["total_validation_checks"] += 1
    if status == "warning":
        metrics["warning_count"] += 1
        metrics[f"warning_{bucket}_count"] += 1
    elif status == "quarantine":
        metrics["quarantine_count"] += 1
        metrics[f"unsupported_{bucket}_count"] += 1


def compute_support_metrics(pred: Dict, cv_text: str) -> Dict:
    metrics = {
        "total_validation_checks": 0,
        "quarantine_count": 0,
        "warning_count": 0,
        "unsupported_skill_count": 0,
        "unsupported_experience_count": 0,
        "unsupported_education_count": 0,
        "unsupported_language_count": 0,
        "unsupported_certification_count": 0,
        "warning_skill_count": 0,
        "warning_experience_count": 0,
        "warning_education_count": 0,
        "warning_language_count": 0,
        "warning_certification_count": 0,
    }

    for skill in _as_list(pred.get("skills")):
        if isinstance(skill, dict):
            value = skill.get("name", "")
            evidence = skill.get("evidence_text", "")
            confidence = _get_confidence(skill)
        else:
            value, evidence, confidence = str(skill), "", None
        _count_status(metrics, _evidence_status(value, evidence, cv_text, confidence), "skill")

    for exp in _as_list(pred.get("experiences")):
        if not isinstance(exp, dict):
            continue
        value = " ".join(
            str(exp.get(key, "")).strip()
            for key in ("role_title", "company_name", "description")
            if exp.get(key)
        )
        evidence = exp.get("evidence_text", "")
        _count_status(metrics, _evidence_status(value, evidence, cv_text, _get_confidence(exp)), "experience")

    for edu in _as_list(pred.get("educations")):
        if not isinstance(edu, dict):
            continue
        value = " ".join(
            str(edu.get(key, "")).strip()
            for key in ("degree", "field_of_study", "institution")
            if edu.get(key)
        )
        _count_status(metrics, _evidence_status(value, edu.get("evidence_text", ""), cv_text, _get_confidence(edu)), "education")

    for language in _as_list(pred.get("languages")):
        value = language.get("name", "") if isinstance(language, dict) else str(language)
        evidence = language.get("evidence_text", "") if isinstance(language, dict) else ""
        confidence = _get_confidence(language) if isinstance(language, dict) else None
        _count_status(metrics, _evidence_status(value, evidence, cv_text, confidence), "language")

    for cert in _as_list(pred.get("certifications")):
        value = cert.get("name", "") if isinstance(cert, dict) else str(cert)
        evidence = cert.get("evidence_text", "") if isinstance(cert, dict) else ""
        confidence = _get_confidence(cert) if isinstance(cert, dict) else None
        _count_status(metrics, _evidence_status(value, evidence, cv_text, confidence), "certification")

    pred_triples = make_triples(pred)
    unsupported_triples = [
        triple for triple in pred_triples
        if not _phrase_supported(triple.object, cv_text)
    ]

    total_checks = metrics["total_validation_checks"]
    total_triples = len(pred_triples)
    metrics.update({
        "unsupported_extraction_rate": round(metrics["quarantine_count"] / total_checks, 3) if total_checks else 0.0,
        "evidence_issue_rate": round((metrics["quarantine_count"] + metrics["warning_count"]) / total_checks, 3) if total_checks else 0.0,
        "total_pred_triples": total_triples,
        "unsupported_triples": len(unsupported_triples),
        "unsupported_triple_rate": round(len(unsupported_triples) / total_triples, 3) if total_triples else 0.0,
    })
    return metrics


def compute_hallucination_rate(pred: Dict, cv_text: str) -> Tuple[float, int, int]:
    metrics = compute_support_metrics(pred, cv_text)
    return (
        metrics["unsupported_extraction_rate"],
        metrics["quarantine_count"],
        metrics["total_validation_checks"],
    )


# ── CV metni okuma ────────────────────────────────────────────────────

def read_cv_text(cv_path: Path) -> str:
    if cv_path.suffix.lower() == ".docx":
        try:
            from docx import Document
            doc = Document(cv_path)
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())
            return "\n".join(parts)
        except Exception as e:
            logger.warning(f"DOCX okuma hatası: {e}")
    elif cv_path.suffix.lower() == ".pdf":
        try:
            import pdfplumber

            with pdfplumber.open(cv_path) as pdf:
                text = "\n\n".join(page.extract_text() or "" for page in pdf.pages).strip()
            if len(text) >= 50:
                return text
        except Exception as e:
            logger.warning(f"PDF pdfplumber okuma hatasi: {e}")

        try:
            import fitz

            doc = fitz.open(str(cv_path))
            text = "\n\n".join(page.get_text("text") or "" for page in doc).strip()
            doc.close()
            if len(text) >= 50:
                return text
        except Exception as e:
            logger.warning(f"PDF PyMuPDF okuma hatasi: {e}")

        try:
            from pdfminer.high_level import extract_text

            text = extract_text(str(cv_path)).strip()
            if len(text) >= 50:
                return text
        except Exception as e:
            logger.warning(f"PDF pdfminer okuma hatasi: {e}")
    elif cv_path.suffix.lower() == ".txt":
        return cv_path.read_text(encoding="utf-8")
    return ""


# ── Model adını kısa tablo etiketine çevir ────────────────────────────

def short_name(model: str) -> str:
    return model.split("/")[-1]


# ── Ana çalıştırıcı ───────────────────────────────────────────────────

CONFIGS = {
    "BL-1": "Zero-shot",
    "BL-2": "Few-shot",
    "SYS-A": "Few-shot + CoT",
    "SYS-B": "Few-shot + CoT + Semantic Inference",
    "SYS-C": "Tam Sistem",
}


def source_file_of(item: Dict[str, Any]) -> str | None:
    """Return the CV file name across old and new gold/result schemas."""
    return item.get("source_file") or item.get("file_name")


def run_ablation(
    cv_dir: str,
    gold_path: str = "evaluation/golden_extractions_pretty.json",
    n_cvs: int = 20,
    configs: List[str] = None,
    models: List[str] = None,
    output_path: str = "evaluation/ablation_results.json",
    resume: bool = True,
):
    cv_dir_path = Path(cv_dir)
    configs = configs or ["SYS-B"]
    models  = models  or ["Qwen/Qwen2.5-7B-Instruct"]

    with open(gold_path, encoding="utf-8") as f:
        gold_data = json.load(f)
    gold_data = gold_data[:n_cvs]
    logger.info(f"✅ {len(gold_data)} CV gold standard yüklendi")

    results: Dict = {}
    if resume and Path(output_path).exists():
        with open(output_path, encoding="utf-8") as f:
            results = json.load(f)
        logger.info(f"📂 Resume: {len(results)} önceki run yüklendi")

    for config_name in configs:
        for model in models:
            run_key = f"{config_name}_{short_name(model)}"
            if run_key not in results:
                results[run_key] = {
                    "config": config_name,
                    "model": model,
                    "cv_results": [],
                    "aggregate": {},
                }

            existing_cvs = {
                source_file_of(r)
                for r in results[run_key]["cv_results"]
                if "error" not in r and source_file_of(r)
            }
            system_prompt, build_user = get_prompt_config(config_name)

            logger.info(f"\n{'='*55}")
            logger.info(f"Config: {config_name} | Model: {model}")
            logger.info(f"{'='*55}")

            for i, gold_item in enumerate(gold_data):
                source_file = source_file_of(gold_item)
                if not source_file:
                    raise KeyError(
                        "Gold item must include either 'source_file' or 'file_name'. "
                        f"Available keys: {sorted(gold_item.keys())}"
                    )
                if source_file in existing_cvs:
                    logger.info(f"  ⏭️  [{i+1}/{len(gold_data)}] {gold_item['candidate_name']} — atlanıyor")
                    continue

                cv_path = cv_dir_path / source_file
                if not cv_path.exists():
                    results[run_key]["cv_results"].append({
                        "source_file": source_file,
                        "candidate_name": gold_item["candidate_name"],
                        "error": "file_missing",
                    })
                    _save(results, output_path)
                    logger.warning(f"  ⚠️  [{i+1}/{len(gold_data)}] Dosya yok: {source_file}")
                    continue

                cv_text = read_cv_text(cv_path)
                if len(cv_text) < 50:
                    results[run_key]["cv_results"].append({
                        "source_file": source_file,
                        "candidate_name": gold_item["candidate_name"],
                        "error": "text_too_short",
                        "text_length": len(cv_text),
                    })
                    _save(results, output_path)
                    logger.warning(f"  ⚠️  [{i+1}/{len(gold_data)}] Metin çok kısa: {source_file}")
                    continue

                logger.info(f"  [{i+1}/{len(gold_data)}] {gold_item['candidate_name']}...")

                user_prompt = build_user(cv_text)

                start_time = time.time()
                pred = call_llm(system_prompt, user_prompt, model)
                elapsed_sec = round(time.time() - start_time, 2)

                if pred is None:
                    logger.warning(f"    ❌ Başarısız ({elapsed_sec}s)")
                    results[run_key]["cv_results"].append({
                        "source_file": source_file,
                        "candidate_name": gold_item["candidate_name"],
                        "error": "LLM failed",
                        "elapsed_sec": elapsed_sec,
                    })
                    _save(results, output_path)
                    continue

                ner  = compute_ner_f1(pred, gold_item)
                re_f1 = compute_re_f1(pred, gold_item)
                kg_triple = compute_kg_triple_metrics(pred, gold_item)
                support_metrics = compute_support_metrics(pred, cv_text)
                hall_rate = support_metrics["unsupported_extraction_rate"]
                hall_count = support_metrics["quarantine_count"]
                total_checks = support_metrics["total_validation_checks"]

                results[run_key]["cv_results"].append({
                    "source_file": source_file,
                    "candidate_name": gold_item["candidate_name"],
                    "domain": gold_item.get("domain", ""),
                    "elapsed_sec": elapsed_sec,
                    "prediction": pred,
                    "ner_f1": ner,
                    "re_f1": re_f1,
                    "kg_triple": kg_triple,
                    "hallucination_rate": hall_rate,
                    "unsupported_count": hall_count,
                    "total_validation_checks": total_checks,
                    "support_metrics": support_metrics,
                    "hallucinated_skills": hall_count,
                    "total_skills": total_checks,
                    "predicted_skill_count": len(_as_list(pred.get("skills"))),
                    "gold_skill_count": len(_as_list(gold_item.get("skills"))),
                })

                logger.info(
                    f"    NER={ner['overall_f1']:.3f} "
                    f"skill={ner['skill_f1']:.3f} "
                    f"RE={re_f1:.3f} "
                    f"KG={kg_triple['f1']:.3f} "
                    f"hall={hall_rate:.1%} "
                    f"utri={support_metrics['unsupported_triple_rate']:.1%} "
                    f"⏱{elapsed_sec}s"
                )

                _save(results, output_path)
                time.sleep(0.3)

            _compute_aggregate(results[run_key])
            _save(results, output_path)

    _print_summary(results)
    _save_summary(results, output_path.replace(".json", "_summary.txt"))
    logger.info(f"\n✅ Tamamlandı → {output_path}")
    return results


# ── Aggregate + raporlama ─────────────────────────────────────────────

def _compute_aggregate(run_data: Dict):
    latest_by_source = {}
    for result in run_data["cv_results"]:
        latest_by_source[result.get("source_file", f"row_{len(latest_by_source)}")] = result
    all_results = list(latest_by_source.values())
    ok = [r for r in all_results if "error" not in r]
    if not ok:
        run_data["aggregate"] = {
            "n_cvs": 0,
            "attempted_cvs": len(all_results),
            "success_rate": 0.0,
            "failure_breakdown": dict(Counter(r.get("error", "success") for r in all_results)),
        }
        return
    n = len(ok)
    attempted = len(all_results)
    failure_breakdown = Counter(r.get("error", "success") for r in all_results)
    kg_tp = sum(r.get("kg_triple", {}).get("tp", 0) for r in ok)
    kg_pred = sum(r.get("kg_triple", {}).get("pred_count", 0) for r in ok)
    kg_gold = sum(r.get("kg_triple", {}).get("gold_count", 0) for r in ok)
    kg_precision = kg_tp / kg_pred if kg_pred else 0.0
    kg_recall = kg_tp / kg_gold if kg_gold else 0.0
    kg_f1 = 0.0 if kg_precision + kg_recall == 0 else 2 * kg_precision * kg_recall / (kg_precision + kg_recall)
    run_data["aggregate"] = {
        "n_cvs": n,
        "attempted_cvs": attempted,
        "success_rate": round(n / attempted, 3) if attempted else 0.0,
        "failure_breakdown": dict(failure_breakdown),
        "avg_skill_f1":          round(sum(r["ner_f1"]["skill_f1"]   for r in ok) / n, 3),
        "avg_company_f1":        round(sum(r["ner_f1"]["company_f1"] for r in ok) / n, 3),
        "avg_education_f1":      round(sum(r["ner_f1"]["education_f1"] for r in ok) / n, 3),
        "avg_overall_ner_f1":    round(sum(r["ner_f1"]["overall_f1"] for r in ok) / n, 3),
        "avg_re_f1":             round(sum(r["re_f1"]               for r in ok) / n, 3),
        "kg_triple_precision":    round(kg_precision, 3),
        "kg_triple_recall":       round(kg_recall, 3),
        "kg_triple_f1":           round(kg_f1, 3),
        "kg_triple_tp":           kg_tp,
        "kg_triple_pred_count":   kg_pred,
        "kg_triple_gold_count":   kg_gold,
        "avg_hallucination_rate":round(sum(r["hallucination_rate"]   for r in ok) / n, 3),
        "avg_evidence_issue_rate": round(sum(r.get("support_metrics", {}).get("evidence_issue_rate", 0.0) for r in ok) / n, 3),
        "avg_unsupported_triple_rate": round(sum(r.get("support_metrics", {}).get("unsupported_triple_rate", 0.0) for r in ok) / n, 3),
        "total_hallucinated":    sum(r.get("unsupported_count", r.get("hallucinated_skills", 0)) for r in ok),
        "total_skills":          sum(r.get("total_validation_checks", r.get("total_skills", 0)) for r in ok),
        "total_unsupported_extractions": sum(r.get("support_metrics", {}).get("quarantine_count", 0) for r in ok),
        "total_warning_extractions": sum(r.get("support_metrics", {}).get("warning_count", 0) for r in ok),
        "total_validation_checks": sum(r.get("support_metrics", {}).get("total_validation_checks", 0) for r in ok),
        "total_unsupported_triples": sum(r.get("support_metrics", {}).get("unsupported_triples", 0) for r in ok),
        "total_pred_triples": sum(r.get("support_metrics", {}).get("total_pred_triples", 0) for r in ok),
        "avg_elapsed_sec":       round(sum(r["elapsed_sec"]   for r in ok) / n, 2),
        "total_elapsed_sec":     round(sum(r["elapsed_sec"]   for r in ok), 1),
    }
    agg = run_data["aggregate"]
    logger.info(
        f"  📊 NER={agg['avg_overall_ner_f1']:.3f} "
        f"RE={agg['avg_re_f1']:.3f} "
        f"KG={agg['kg_triple_f1']:.3f} "
        f"Hall={agg['avg_hallucination_rate']:.1%} "
        f"UTriple={agg['avg_unsupported_triple_rate']:.1%} "
        f"success={agg['success_rate']:.1%} "
        f"⏱avg={agg['avg_elapsed_sec']}s  n={n}"
    )


def _save(results: Dict, path: str):
    output_path = Path(path)
    output_path.parent.mkdir(exist_ok=True, parents=True)
    tmp_path = output_path.with_name(f"{output_path.name}.tmp")
    last_error = None
    for attempt in range(5):
        try:
            with open(tmp_path, "w", encoding="utf-8") as f:
                json.dump(results, f, ensure_ascii=False, indent=2)
            os.replace(tmp_path, output_path)
            return
        except OSError as e:
            last_error = e
            logger.warning(f"Kaydetme tekrar denenecek ({attempt + 1}/5): {e}")
            time.sleep(0.5 * (attempt + 1))
    raise last_error


def _print_summary(results: Dict):
    print("\n" + "=" * 100)
    print("  ABLATION / MODEL KARŞILAŞTIRMA SONUÇLARI")
    print("=" * 100)
    print(f"  {'Run Key':<38} {'NER F1':>7} {'Skill':>6} {'RE F1':>6} {'KG F1':>6} {'Hall%':>6} {'UTr%':>6} {'Succ%':>6} {'t(s)':>6} {'N':>3}")
    print(f"  {'-'*38} {'-'*7} {'-'*6} {'-'*6} {'-'*6} {'-'*6} {'-'*6} {'-'*6} {'-'*6} {'-'*3}")
    for key, data in sorted(results.items()):
        agg = data.get("aggregate", {})
        if not agg:
            continue
        print(
            f"  {key:<38} "
            f"{agg.get('avg_overall_ner_f1',0):>7.3f} "
            f"{agg.get('avg_skill_f1',0):>6.3f} "
            f"{agg.get('avg_re_f1',0):>6.3f} "
            f"{agg.get('kg_triple_f1',0):>6.3f} "
            f"{agg.get('avg_hallucination_rate',0):>5.1%} "
            f"{agg.get('avg_unsupported_triple_rate',0):>5.1%} "
            f"{agg.get('success_rate',0):>5.1%} "
            f"{agg.get('avg_elapsed_sec',0):>6.1f} "
            f"{agg.get('n_cvs',0):>3}"
        )
    print("=" * 100)


def _save_summary(results: Dict, path: str):
    lines = [
        "ABLATION / MODEL KARŞILAŞTIRMA SONUÇLARI\n",
        "=" * 100 + "\n",
        f"{'Run Key':<38} {'NER F1':>7} {'Skill':>6} {'RE F1':>6} {'KG F1':>6} {'Hall%':>6} {'UTr%':>6} {'Succ%':>6} {'t(s)':>6} {'N':>3}\n",
        f"{'-'*38} {'-'*7} {'-'*6} {'-'*6} {'-'*6} {'-'*6} {'-'*6} {'-'*6} {'-'*6} {'-'*3}\n",
    ]
    for key, data in sorted(results.items()):
        agg = data.get("aggregate", {})
        if not agg:
            continue
        lines.append(
            f"{key:<38} "
            f"{agg.get('avg_overall_ner_f1',0):>7.3f} "
            f"{agg.get('avg_skill_f1',0):>6.3f} "
            f"{agg.get('avg_re_f1',0):>6.3f} "
            f"{agg.get('kg_triple_f1',0):>6.3f} "
            f"{agg.get('avg_hallucination_rate',0):>5.1%} "
            f"{agg.get('avg_unsupported_triple_rate',0):>5.1%} "
            f"{agg.get('success_rate',0):>5.1%} "
            f"{agg.get('avg_elapsed_sec',0):>6.1f} "
            f"{agg.get('n_cvs',0):>3}\n"
        )
    Path(path).parent.mkdir(exist_ok=True, parents=True)
    with open(path, "w", encoding="utf-8") as f:
        f.writelines(lines)
    logger.info(f"📄 Özet → {path}")


# ── CLI ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--cv_dir",  default="data/cvs")
    parser.add_argument("--gold",    default="evaluation/golden_extractions_pretty.json")
    parser.add_argument("--n_cvs",   type=int, default=20)
    parser.add_argument("--configs", nargs="+", default=["SYS-B"], choices=list(CONFIGS.keys()))
    parser.add_argument("--models",  nargs="+", default=["Qwen/Qwen2.5-7B-Instruct"])
    parser.add_argument("--output",  default="evaluation/ablation_results.json")
    parser.add_argument("--no_resume", action="store_true")
    args = parser.parse_args()

    logger.info("=" * 60)
    logger.info("TalentForge Ablation Study")
    logger.info(f"CV: {args.n_cvs} | Configs: {args.configs}")
    logger.info(f"Models: {args.models}")
    logger.info("=" * 60)

    run_ablation(
        cv_dir=args.cv_dir,
        gold_path=args.gold,
        n_cvs=args.n_cvs,
        configs=args.configs,
        models=args.models,
        output_path=args.output,
        resume=not args.no_resume,
    )
